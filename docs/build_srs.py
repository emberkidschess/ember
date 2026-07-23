from __future__ import annotations

import hashlib
import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
EVIDENCE = ROOT / "docs" / "evidence"
OUT = ROOT / "docs" / "EmberKids_SRS_Project_Handover.docx"

NAVY = "111111"
BLUE = "222222"
MID_BLUE = "333333"
ORANGE = "111111"
INK = "202124"
MUTED = "666666"
LIGHT_BLUE = "F1F1F1"
LIGHT_ORANGE = "F7F7F7"
WHITE = "FFFFFF"
GRID = "C8C8C8"
GREEN = "F4F4F4"
YELLOW = "FAFAFA"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w * 1440)))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(i, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_run(run, size=11, color=INK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def style_paragraph(p, before=0, after=6, line=1.25, keep=False):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep:
        fmt.keep_with_next = True


def add_text(doc, text, style=None, bold_prefix=None, color=INK, after=6, before=0, italic=False):
    p = doc.add_paragraph(style=style)
    style_paragraph(p, before=before, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, bold=True, color=NAVY)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_run(r, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    style_paragraph(p, after=4, line=1.2)
    r = p.add_run(text)
    set_run(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.23)
    style_paragraph(p, after=4, line=1.2)
    r = p.add_run(text)
    set_run(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), color={1: BLUE, 2: BLUE, 3: MID_BLUE}.get(level, NAVY), bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = False
    style_paragraph(p, before=2, after=10, line=1.0)
    r = p.add_run(text)
    set_run(r, size=9, color=MUTED, italic=True)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.0)
        r = p.add_run(text)
        set_run(r, size=font_size, color=NAVY, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if row_index % 2 == 1:
                set_cell_shading(cells[i], "F7F9FC")
            p = cells[i].paragraphs[0]
            style_paragraph(p, after=0, line=1.1)
            r = p.add_run(str(value))
            set_run(r, size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_callout(doc, title, body, fill=LIGHT_ORANGE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [6.5])
    set_table_borders(table, color="8A8A8A", size="8")
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, after=3, line=1.15)
    r = p.add_run(title)
    set_run(r, color=ORANGE, bold=True, size=10)
    p2 = cell.add_paragraph()
    style_paragraph(p2, after=0, line=1.2)
    r2 = p2.add_run(body)
    set_run(r2, color=INK, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_picture_with_caption(doc, path: Path, caption: str, width=6.25):
    if not path.exists():
        add_callout(doc, "Evidence image unavailable", f"Expected file: {path.name}", fill=YELLOW)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = False
    inline = p.add_run().add_picture(str(path), width=Inches(width))
    # Set accessible alt text on the embedded picture for screen readers.
    descr = inline._inline.docPr
    descr.set("title", caption)
    descr.set("descr", caption)
    add_caption(doc, caption)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, color, before, after in ((1, 16, BLUE, 14, 8), (2, 13, BLUE, 11, 6), (3, 12, MID_BLUE, 8, 4)):
        s = styles[f"Heading {level}"]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    if "Evidence Caption" not in styles:
        s = styles.add_style("Evidence Caption", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = "Calibri"
        s.font.size = Pt(9)
        s.font.italic = True
        s.font.color.rgb = RGBColor.from_string(MUTED)


def configure_section(section):
    section.top_margin = Inches(0.92)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(hp, after=0, line=1.0)
    r = hp.add_run("EMBERKIDS  /  CONFIDENTIAL PROJECT RECORD")
    set_run(r, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    style_paragraph(fp, after=0, line=1.0)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("EmberKids Chess Academy  •  SRS & Handover  •  ")
    set_run(r, size=8.5, color=MUTED)
    add_field(fp, "PAGE")


def title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, before=18, after=4, line=1.0)
    r = p.add_run("EMBERKIDS")
    set_run(r, size=13, color=ORANGE, bold=True)
    r2 = p.add_run("  /  PROJECT DELIVERY DOCUMENT")
    set_run(r2, size=9, color=MUTED, bold=True)

    p = doc.add_paragraph()
    style_paragraph(p, before=6, after=3, line=0.98, keep=True)
    r = p.add_run("Software Requirements Specification")
    set_run(r, size=25, color=NAVY, bold=True)
    p2 = doc.add_paragraph()
    style_paragraph(p2, after=14, line=1.0, keep=True)
    r = p2.add_run("& Project Handover Record")
    set_run(r, size=20, color=ORANGE, bold=True)

    add_text(doc, "A complete scope, acceptance, evidence and commercial handover record for the EmberKids Chess Academy platform.", color=MUTED, after=14, italic=True)
    metadata = [
        ["Project", "EmberKids Chess Academy Platform"],
        ["Document owner", "[Developer Name]"],
        ["Client / academy", "[Client Legal Name] / EmberKids Chess Academy"],
        ["Version / date", "1.0  •  21 July 2026"],
        ["Status", "Draft for scope confirmation and handover"],
        ["Confidentiality", "Confidential — commercial and technical record"],
    ]
    add_table(doc, ["Field", "Record"], metadata, [1.65, 4.85], font_size=10)
    add_callout(doc, "Before sending", "Replace every bracketed placeholder, insert the agreed amount and payment dates, attach the final repository/archive hash, and have a lawyer review the ownership and payment clauses. This document records the project; it is not legal advice.")


def build():
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    title_block(doc)

    add_page_break(doc)
    add_heading(doc, "1. Purpose, scope and evidentiary use", 1)
    add_text(doc, "This document defines the implemented scope of the EmberKids Chess Academy website and academy-management platform, records the current handover expectations, and provides a dated evidence register. It is intended to be signed by the developer and client so both parties have a shared record of what was delivered, what remains client-owned or client-supplied, and how payment and maintenance are handled.")
    add_text(doc, "The document is written as a project record, not as a replacement for a lawyer-drafted development agreement. Where a clause affects ownership, licensing, payment recovery or enforceability, the parties should obtain legal review under the applicable Indian law and identify the correct legal entity and signatory.", color=MUTED, italic=True)
    add_heading(doc, "1.1 Document conventions", 2)
    add_bullet(doc, "Implemented / in scope means the capability is represented in the current codebase and should be verified jointly during acceptance.")
    add_bullet(doc, "Existing-feature bug means a reproducible defect in a feature listed in this SRS, using supported environments, without a new requirement or unauthorized code change.")
    add_bullet(doc, "Client-provided means the client supplies the account, credentials, content, approvals, hosting, domain or third-party subscription needed for production operation.")
    add_bullet(doc, "Bracketed values such as [Final Amount] are fill-in fields and must be completed before signature.")

    add_heading(doc, "2. Product overview", 1)
    add_text(doc, "EmberKids is a responsive chess-academy platform combining a public marketing site, secure admin operations, staff/coach workflows, student self-service, enrolment and payment operations, class and attendance tracking, evaluation/report cards, notifications, and an academy-focused AI guide named Amber.")
    add_heading(doc, "2.1 Solution architecture recorded in the codebase", 2)
    arch_rows = [
        ["Frontend", "Next.js / React / TypeScript responsive web application with public, admin, staff and student routes."],
        ["Backend", "Express / Node.js REST API with controllers, services, validation, scheduled jobs and integrations."],
        ["Data", "MongoDB models for academy, learner, staff, course, batch, attendance, payments, notifications, audit and knowledge data."],
        ["Infrastructure", "Redis-compatible caching is supported/recommended; production requires secure environment configuration, SMTP and a MongoDB replica set for transactions."],
        ["Integrations", "Email/SMTP, WhatsApp messaging, Cloudinary/file storage, payment links/Wise flows and Gemini Flash + embeddings for the academy chatbot."],
    ]
    add_table(doc, ["Layer", "Recorded implementation"], arch_rows, [1.25, 5.25], font_size=9.4)
    add_callout(doc, "Configuration note", "The AI model is selected server-side through configuration. The original project scope called for Gemini 2.5 Flash; the deployed model name should be confirmed from the production environment before final acceptance. Never place an API key, password or private token in this SRS.", fill=YELLOW)

    add_heading(doc, "3. Implemented feature inventory", 1)
    add_text(doc, "The following inventory is grouped for the client call and for acceptance testing. Each group should be checked against the live deployment before the final sign-off.")
    features = [
        ["Public site", "Home/landing page, academy positioning, calls-to-action, about, courses and roadmap, coaches, prodigies, testimonials, FAQ, contact, privacy and terms pages; responsive layouts, theme styling, animation, forms and contact links.", "In scope"],
        ["Admissions and leads", "Free-trial/admission enquiry capture, lead records, follow-up workflow, contact details, demo-class journey and lead status management.", "In scope"],
        ["Admin portal", "Secure admin login/password reset, dashboard, role/permission controls, students, leads, courses/roadmap, testimonials, site content, coaches/staff, batches/history, classes, trial classes, events/masterclasses, tournaments, packages, payment links, payments, exports, reports and audit logs.", "In scope"],
        ["Staff / coach portal", "Staff login, dashboard, assigned classes and batches, student lists, attendance and disputes, trial classes, report cards/evaluations, reports and payment-link/salary-related workflows represented by the current modules.", "In scope"],
        ["Student portal", "Student login/password reset, dashboard, upcoming classes, schedule and batch information, attendance, evaluation/report cards, packages/payments, notifications, WhatsApp group link, profile and account flows.", "In scope"],
        ["Core API and security", "REST routes, controllers/services, MongoDB models, refresh-token/session handling, role-based permissions, secure cookies, CORS/Helmet/sanitisation, rate limits, audit logs, transactions, scheduled jobs, email, WhatsApp, uploads/storage and CSV exports.", "In scope"],
        ["Amber AI guide", "Gemini Flash chat integration, website/content knowledge chunks, embeddings/indexing, retrieval-grounded prompts, academy-only boundary, out-of-scope fallback with consultant contact, fee enquiries routed to the academy team, clear batch-format explanation, streaming responses, session history, typing/loading/error states, responsive popup, Amber mascot launcher and greeting/tap audio.", "In scope"],
    ]
    add_table(doc, ["Area", "Included capability", "Status"], features, [1.25, 4.55, 0.70], font_size=8.7)

    add_heading(doc, "3.1 Academy-specific chatbot behavior", 2)
    add_bullet(doc, "Amber answers academy-related questions using retrieved academy knowledge first, including courses, batches, fees, FAQs, instructors, policies and contact details available in the indexed content.")
    add_bullet(doc, "For unrelated questions, Amber politely states that it can assist only with EmberKids academy information and offers the academy contact route.")
    add_bullet(doc, "For fee questions, Amber directs the visitor to the academy consultant/team for the current fee conversation and provides configured contact details rather than inventing a price.")
    add_bullet(doc, "For batch-size questions, Amber explains the three formats clearly: 1:1, Premium Group (2–3 students), and Standard Group (5–6 students), subject to the academy’s current availability.")
    add_bullet(doc, "The launcher is designed to sit near the WhatsApp action without overlap; the tooltip and message behavior are part of the UI acceptance check.")

    add_heading(doc, "4. Functional requirements and acceptance matrix", 1)
    fr_rows = [
        ["FR-01", "Public navigation", "All public routes load, navigation and responsive layout work, and key CTA/contact links resolve.", "Joint UAT"],
        ["FR-02", "Enquiry / trial capture", "A visitor can submit a supported enquiry or trial request; validation and lead creation are visible to authorised staff/admin.", "Joint UAT"],
        ["FR-03", "Admin operations", "Authorised admin can manage the listed academy records, content, batches, schedules, payments and reports.", "Joint UAT"],
        ["FR-04", "Staff workflows", "Authorised staff can view assigned work, attendance, learners, trial classes and evaluation/report modules allowed by permissions.", "Joint UAT"],
        ["FR-05", "Student self-service", "A student can authenticate and view their schedule, attendance, report cards, payments/packages and notifications as configured.", "Joint UAT"],
        ["FR-06", "Payments and links", "Configured payment links and status callbacks create the expected payment/enrolment state; provider availability is a dependency.", "Joint UAT"],
        ["FR-07", "Notifications", "Configured email/WhatsApp/in-app notifications are generated for supported events and failures are observable in logs/audit records.", "Joint UAT"],
        ["FR-08", "Amber RAG chat", "Every query retrieves relevant knowledge before generation, streaming works when provider is available, and out-of-scope/fee/batch rules are respected.", "Joint UAT"],
        ["FR-09", "Chat UX", "Launcher, popup, typing indicator, auto-scroll, Enter/Shift+Enter, retry/error state, current-session history, responsive width/height and no-overlap placement work on supported breakpoints.", "Joint UAT"],
        ["FR-10", "Audit and export", "Authorised users can access supported audit records and CSV/report exports without exposing secrets.", "Joint UAT"],
    ]
    add_table(doc, ["ID", "Requirement", "Acceptance criterion", "Verification"], fr_rows, [0.62, 1.35, 3.75, 0.78], font_size=8.5)

    add_heading(doc, "5. Non-functional requirements", 1)
    nfr_rows = [
        ["Responsive UX", "Public, portal and chatbot layouts adapt to desktop, tablet and small screens without overlap or unusable controls."],
        ["Performance", "Keep client bundles, network requests, database queries and RAG context focused on the current task; exact production targets depend on hosting and third-party latency."],
        ["Security", "Use server-side secrets, secure cookies/tokens, access control, validation/sanitisation, rate limiting, security headers, least privilege and audit trails."],
        ["Availability", "Third-party services (Gemini, SMTP, WhatsApp, payment providers, storage and hosting) may be unavailable; the UI must show graceful error/retry states."],
        ["Accessibility", "Use keyboard-accessible controls, visible focus, labels, sensible contrast and reduced-motion-friendly behavior where supported."],
        ["Privacy", "Do not store or publish secrets in source, screenshots, SRS or logs. Collect only the learner/contact data needed for academy operations and apply the client’s privacy policy."],
        ["Maintainability", "Keep environment variables documented, modules separated by responsibility, migrations/backups planned, and changes recorded in Git and the change log."],
    ]
    add_table(doc, ["Quality area", "Requirement"], nfr_rows, [1.55, 4.95], font_size=9.1)

    add_heading(doc, "6. Assumptions, dependencies and exclusions", 1)
    add_heading(doc, "6.1 Client responsibilities and dependencies", 2)
    for text in [
        "The client supplies and pays for hosting, domain, DNS/SSL, database/Redis services, SMTP, Gemini/API, WhatsApp, Cloudinary/storage, payment-provider and other third-party accounts or usage charges.",
        "The client supplies accurate academy content, pricing/fee approvals, batch availability, contact numbers, policies, logos/assets, legal text and production credentials in a secure channel.",
        "The client approves UAT findings, signs the acceptance record, maintains backups and controls production account access after handover.",
        "Provider/model changes, quotas, rate limits, account suspension, network outages or changes made by third parties are dependencies and may require a separate change request.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "6.2 Exclusions unless separately agreed", 2)
    for text in [
        "New features, new portal roles, new pages, new integrations, major redesigns, rebranding, new reports or changes to business rules after acceptance.",
        "Bulk content entry, data cleansing/migration, ongoing academy operations, learner support, 24/7 monitoring, hosting/domain fees, third-party subscriptions and provider usage fees.",
        "Security certification, legal/compliance drafting, penetration testing, SEO/marketing campaigns or guarantees of third-party ranking/deliverability.",
        "Fixes caused by unauthorised code changes, unsupported environments, client infrastructure misconfiguration, expired credentials or third-party outages.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "7. Handover, payment and intellectual-property record", 1)
    add_callout(doc, "Important", "A repository or hosting account being created in the client’s name does not, by itself, prove that source-code ownership or payment was transferred. The parties should sign this record (or a separate development agreement), keep the payment trail, and create a dated source archive/hash before final handover.", fill=YELLOW)
    add_heading(doc, "7.1 Proposed handover package", 2)
    handover_rows = [
        ["1", "Signed SRS and acceptance record", "Scope, exclusions, acceptance notes and signatures."],
        ["2", "Source repository / archive", "Source code, Git history, build files and a dated SHA-256 hash."],
        ["3", "Deployment pack", "README, environment-variable template, build/start instructions and rollback notes; no secret values."],
        ["4", "Data and integration notes", "Models/routes/integrations, provider assumptions, backup expectations and account ownership list."],
        ["5", "Evidence pack", "Dated screenshots, commit log, change log and test/UAT notes."],
    ]
    add_table(doc, ["#", "Item", "Record"], handover_rows, [0.42, 1.75, 4.33], font_size=9.0)
    add_heading(doc, "7.2 Payment and rights clause to complete and sign", 2)
    add_text(doc, "The commercial fields below are intentionally blank until the parties enter the agreed values:")
    payment_rows = [
        ["Agreed project amount", "[₹__________]", "Final amount including/excluding taxes: [__________]"],
        ["Advance / milestone 1", "[₹__________] due [date]", "Trigger: [__________]"],
        ["Milestone 2", "[₹__________] due [date]", "Trigger: [__________]"],
        ["Final payment", "[₹__________] due [date]", "Due before full source/IP handover"],
        ["Outstanding at signature", "[₹__________]", "To be acknowledged by both parties"],
    ]
    add_table(doc, ["Commercial field", "Value", "Notes"], payment_rows, [1.55, 1.65, 3.30], font_size=9.0)
    add_number(doc, "Until the agreed amount is paid in full and cleared, the developer retains ownership of the source code and deliverables, and the client receives only the limited evaluation/use permission expressly agreed in writing.")
    add_number(doc, "After full cleared payment and written handover, the developer transfers or licenses the agreed rights to the client as stated in the signed commercial agreement. The exact legal form of the transfer should be confirmed by counsel.")
    add_number(doc, "Before full payment and written handover, the client must not copy, resell, sublicense, commercialise, publish, modify for another provider, or distribute the source code except as expressly agreed in writing.")
    add_number(doc, "No API key, password, private token, database credential or personal data is part of the SRS. Credentials must be exchanged separately and rotated at handover.")
    add_number(doc, "If the client requests work beyond the listed scope, the parties should approve a written change request with effort, price and delivery effect before that work begins.")

    add_heading(doc, "8. Three-month maintenance policy", 1)
    add_callout(doc, "Promised support window", "The developer is providing three (3) months of maintenance after [written acceptance / production launch — choose one], subject to the limits below. This is a limited bug-fix window for existing implemented features, not an open-ended development retainer.", fill=GREEN)
    maintenance_rows = [
        ["Included", "Reproducible bugs in features listed in this SRS that were already implemented and accepted, when used with the agreed production configuration and supported browsers/devices."],
        ["Included", "Reasonable diagnosis, a corrective code change, and a verification/deployment note for the existing feature where the developer still has authorised access."],
        ["Response", "Client reports should include page/feature, steps to reproduce, expected result, actual result, timestamp and screenshots/video where useful. The developer will acknowledge and prioritise reasonably; no 24/7 SLA is promised unless separately signed."],
        ["Not included", "New features, change requests, redesigns, content changes, new integrations, new reports, new roles, data entry/migration or changes to business rules."],
        ["Not included", "Hosting/domain/DNS/SSL, expired credentials, third-party provider outages or API/model/quota changes, payment/WhatsApp/SMTP/storage issues outside the developer’s code, or client-side misconfiguration."],
        ["Not included", "Defects caused by unauthorised edits, another vendor, unsupported browser/device, malicious activity, force majeure or a production environment that differs materially from the accepted configuration."],
        ["After window", "Requests after the three-month window, or excluded work during the window, require a separate quote or maintenance agreement."],
    ]
    add_table(doc, ["Policy item", "Definition"], maintenance_rows, [1.35, 5.15], font_size=8.9)
    add_text(doc, "Maintenance start date: [__________]   •   Maintenance end date: [__________]   •   Support channel: [__________]", color=MUTED, italic=True)

    add_heading(doc, "9. Acceptance, sign-off and change control", 1)
    add_text(doc, "The parties should complete a joint UAT review of the listed acceptance criteria. The client should provide any defect list in writing with reproduction details during the agreed review period. New requests or preference changes are not defects and should be logged as change requests.")
    acceptance_rows = [
        ["Review period", "[__] calendar days from delivery / access"],
        ["Accepted build / URL", "[Production URL or release tag]"],
        ["Open defects", "[None / list ticket IDs in attachment]"],
        ["Acceptance date", "[DD/MM/YYYY]"],
        ["Payment due on acceptance", "[₹__________] due [date]"],
    ]
    add_table(doc, ["Acceptance field", "Value"], acceptance_rows, [2.05, 4.45], font_size=9.2)
    sign_rows = [
        ["Developer / project author", "[Name]", "Signature: __________________", "Date: __________"],
        ["Client / authorised signatory", "[Legal name + role]", "Signature: __________________", "Date: __________"],
    ]
    add_table(doc, ["Party", "Name", "Signature", "Date"], sign_rows, [1.65, 2.05, 1.90, 0.90], font_size=8.9)

    add_page_break(doc)
    add_heading(doc, "Appendix A — Dated UI evidence", 1)
    add_text(doc, "The following screenshots were captured from the local preview on 21 July 2026. They show the visible website implementation only; confirm the hosted build and current data before signing. No secrets are included.", color=MUTED, italic=True)
    add_picture_with_caption(doc, EVIDENCE / "homepage.png", "Figure A-1. Public homepage viewport — brand, navigation, hero CTA and floating actions.", width=6.2)
    add_picture_with_caption(doc, EVIDENCE / "courses.png", "Figure A-2. Public courses page viewport — course/roadmap presentation and responsive site styling.", width=6.2)
    add_picture_with_caption(doc, EVIDENCE / "contact.png", "Figure A-3. Public contact page viewport — enquiry/contact route and academy-facing information.", width=6.2)
    add_picture_with_caption(doc, EVIDENCE / "chatbot-popup.png", "Figure A-4. Amber chatbot popup — launcher interaction, academy greeting, chat input and responsive visual treatment.", width=6.2)

    add_page_break(doc)
    add_heading(doc, "Appendix B — Evidence register and authorship record", 1)
    evidence_rows = [
        ["A-1", "homepage.png", "Public homepage and floating actions", "Local preview", "21 Jul 2026"],
        ["A-2", "courses.png", "Courses/roadmap presentation", "Local preview", "21 Jul 2026"],
        ["A-3", "contact.png", "Academy contact route", "Local preview", "21 Jul 2026"],
        ["A-4", "chatbot-popup.png", "Amber chatbot popup and launcher", "Local preview", "21 Jul 2026"],
    ]
    add_table(doc, ["ID", "File", "What it records", "Source", "Captured"], evidence_rows, [0.52, 1.35, 2.75, 1.15, 0.73], font_size=8.6)
    add_heading(doc, "B.1 Repository snapshot to attach before handover", 2)
    repo_rows = [
        ["Current branch at evidence review", "main", "Create a signed release tag for final handover."],
        ["Recent commit", "e67878a", "Update chatbot components and add new audio file (19 Jul 2026)."],
        ["Previous chatbot commit", "88820b4", "Add chatbot feature with knowledge base and UI components (19 Jul 2026)."],
        ["Tracked code footprint", "42,863 lines", "TS/TSX/JS/JSX/CSS in frontend and backend at evidence review."],
        ["Module counts", "51 components • 14 services • 28 models • 29 routes", "Use Git export to preserve the exact final snapshot."],
        ["Final archive hash", "[SHA-256: __________________________]", "Compute after final code freeze; store with this signed record."],
    ]
    add_table(doc, ["Evidence item", "Value", "Action / interpretation"], repo_rows, [1.80, 2.05, 2.65], font_size=8.8)
    add_callout(doc, "Authorship protection checklist", "Before handover: commit all final changes, create a signed Git tag, export `git log` and `git status`, create a source archive, compute SHA-256, email this SRS and the hash from the developer’s own account, obtain written client acknowledgement, and keep the payment invoices/receipts. These steps create a stronger contemporaneous record than a screenshot alone.", fill=LIGHT_BLUE)

    add_heading(doc, "Appendix C — Change log", 1)
    change_rows = [
        ["1.0", "21 Jul 2026", "[Developer Name]", "Initial SRS, feature inventory, acceptance matrix, handover/payment clauses, 3-month maintenance policy and dated UI evidence."],
        ["[ ]", "[date]", "[name]", "[Future approved change]"],
    ]
    add_table(doc, ["Version", "Date", "Author", "Change"], change_rows, [0.65, 1.0, 1.25, 3.60], font_size=8.8)

    add_heading(doc, "Appendix D — Final handover acknowledgement", 1)
    add_text(doc, "By signing below, the parties confirm that they have reviewed this SRS, understand the listed scope and exclusions, and will use the agreed commercial agreement/payment record to determine when source-code/IP rights and full handover become effective.")
    ack_rows = [
        ["Client acknowledgement", "I confirm receipt of the SRS and the listed scope/evidence. Outstanding payment and acceptance fields are recorded above.", "Name/signature/date: ______________________________"],
        ["Developer acknowledgement", "I confirm that the listed implementation and maintenance promise reflect the agreed project record, subject to final UAT and the signed commercial terms.", "Name/signature/date: ______________________________"],
    ]
    add_table(doc, ["Party", "Acknowledgement", "Signature"], ack_rows, [1.45, 3.55, 1.50], font_size=8.8)
    add_callout(doc, "Final reminder", "Do not place production API keys, passwords or private client data in this document. Complete the bracketed fields, sign every page or initial the change/evidence pages as appropriate, and obtain legal review for the payment/IP language before relying on it in a dispute.", fill=YELLOW)

    # Core properties are deliberately generic to avoid leaking personal metadata.
    props = doc.core_properties
    props.title = "EmberKids Chess Academy — SRS & Project Handover Record"
    props.subject = "Scope, acceptance, evidence, payment and maintenance record"
    props.author = "Project Developer"
    props.keywords = "EmberKids, SRS, handover, acceptance, maintenance"
    props.comments = "Replace bracketed fields before signature."

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build()
